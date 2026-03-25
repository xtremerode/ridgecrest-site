"""
Generates Ridgecrest_Agent_Architecture.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---------------------------------------------------------------------------
# Styles
# ---------------------------------------------------------------------------
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10)

def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)

def heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return p

def body(text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x2C, 0x2C, 0x54)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1A1A2E')
        tcPr.append(shd)
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Data rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            cell = row.cells[i]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(table, i, w)

    doc.add_paragraph()
    return table

# ---------------------------------------------------------------------------
# Title
# ---------------------------------------------------------------------------
title = doc.add_heading('Ridgecrest Designs — Ad Management Agent Architecture', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph('Full stack reference for Lovable dashboard integration')
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].font.size = Pt(10)
sub.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 1 — Agent Inventory
# ---------------------------------------------------------------------------
heading('1. Agent Inventory')
body('Every agent name as stored in agent_heartbeats.agent_name with its responsibility.')
doc.add_paragraph()

add_table(
    headers=['agent_name', 'File', 'Responsibility'],
    rows=[
        ['orchestrator',        'orchestrator.py',          'Master controller — triggers all agents in sequence on a daily schedule and manages pipeline state'],
        ['google_sync',         'google_sync.py',           'Pulls Google Ads campaign, ad group, keyword, and daily performance data into the local DB'],
        ['meta_sync',           'meta_sync.py',             'Pulls Meta campaign, ad set, and daily insight data into the local DB'],
        ['microsoft_sync',      'microsoft_sync.py',        'Pulls Microsoft Ads campaign, ad group, and daily performance data into the local DB'],
        ['performance_analyst', 'performance_analyst.py',   'Reads DB metrics, evaluates campaign health, writes budget snapshots, triggers alerts'],
        ['meta_manager',        'meta_manager.py',          'Reads Meta performance from DB, enforces active days, pauses/adjusts budgets on Meta API'],
        ['microsoft_manager',   'microsoft_manager.py',     'Activates priority Microsoft campaigns, adjusts bids, enforces day scheduling via Microsoft API'],
        ['google_ads_scheduler','google_ads_scheduler.py',  'Applies Fri/Sat/Sun/Mon AdSchedule criteria to Google Ads campaigns (staged — awaiting developer token)'],
        ['guardrails',          'guardrails.py',            'Validates all proposed actions against spend, CPL, and API health thresholds before execution'],
        ['bid_budget_optimizer','bid_budget_optimizer.py',  'Executes keyword bid changes and campaign budget reallocation on Google Ads API'],
        ['creative_agent',      'creative_agent.py',        'Generates ad copy briefs and refreshes underperforming ads via Claude API'],
        ['reporting_agent',     'reporting_agent.py',       'Writes daily and weekly performance summaries to the reports table'],
        ['supabase_sync',       'supabase_sync.py',         'Pushes all local DB tables to the Lovable Supabase endpoint after every pipeline run'],
    ],
    col_widths=[1.5, 1.7, 3.4],
)

# ---------------------------------------------------------------------------
# Section 2 — Pipeline Execution Order
# ---------------------------------------------------------------------------
heading('2. Pipeline Execution Order & Data Flow')
body('TRIGGER: Daily at 08:00  (schedule.every().day.at("08:00"))')
doc.add_paragraph()

pipeline_lines = [
    'Phase 0 ── Platform Sync (parallel)',
    '           google_sync      → reads Google Ads API → writes campaigns, performance_metrics',
    '           meta_sync        → reads Meta API       → writes campaigns, ad_groups, performance_metrics',
    '           microsoft_sync   → reads Microsoft API  → writes campaigns, ad_groups, performance_metrics',
    '',
    'Phase 1 ── Performance Analysis (always runs, every day)',
    '           performance_analyst reads performance_metrics, campaigns, ad_groups',
    '           → writes budget_snapshots, search_terms',
    '           → sends performance_analysis_complete to orchestrator',
    '           → sends critical_alert if thresholds breached',
    '',
    'Phase 1b ── Platform Managers (always runs, even if guardrails halt Phase 2)',
    '            meta_manager',
    '              Step 0: enforce active days — pause campaigns Tue/Wed/Thu, resume Fri/Sat/Sun/Mon',
    '              Step 1: evaluate ad set CPL → pause or adjust budgets',
    '              Step 2: evaluate campaign CPL → pause if over ceiling',
    '            microsoft_manager',
    '              Step 0: apply DayTimeCriterion (-100% bid on Tue/Wed/Thu)',
    '              Step 1: activate priority campaigns',
    '              Step 2: optimize ad group bids',
    '            google_ads_scheduler  (staged — activates when developer token approved)',
    '',
    '            ↓ guardrails.check_pipeline_state() called here ↓',
    '            If violations found:',
    '              → writes to guardrail_violations',
    '              → orchestrator sets optimization_halted = True',
    '              → Phases 2 & 3 SKIPPED',
    '              → escalation email sent to henry@ridgecrestdesigns.com',
    '',
    'Phase 2 ── Bid & Budget Optimization  (SKIPPED if guardrails halt)',
    '           guardrails.filter_actions(proposed_actions) → allowed / blocked',
    '           bid_budget_optimizer.run(allowed_actions)',
    '           → adjusts Google Ads keyword bids',
    '           → reallocates campaign budgets (hard cap: $125/day total)',
    '',
    'Phase 3 ── Creative Refresh  (SKIPPED on inactive days OR if guardrails halt)',
    '           creative_agent reads ads, creative_briefs, performance_metrics',
    '           → generates new copy via Claude API',
    '           → writes to creative_briefs, ads',
    '',
    'Phase 4 ── Reporting (always runs)',
    '           reporting_agent reads performance_metrics, optimization_actions,',
    '                                   budget_snapshots, agent_messages, creative_briefs',
    '           → writes daily summary to reports',
    '',
    'Phase 4b ── Weekly Report (Monday only)',
    '            reporting_agent extended run → writes weekly summary to reports',
    '',
    'Phase 5 ── Supabase Sync (always runs last)',
    '           supabase_sync reads ALL tables',
    '           → POSTs to SUPABASE_INGEST_ENDPOINT',
    '           → Lovable dashboard receives fresh data',
]
for line in pipeline_lines:
    code_block(line)

doc.add_paragraph()
body('Background daemon loops:')
code_block('schedule.every().hour        → print_system_status')
code_block('schedule.every(5).minutes   → process_messages()  (drains agent_messages queue)')
doc.add_paragraph()

# ---------------------------------------------------------------------------
# Section 3 — Agent-to-Table Mapping
# ---------------------------------------------------------------------------
heading('3. Agent-to-Table Mapping')
body('Which database tables each agent reads from and writes to.')
doc.add_paragraph()

add_table(
    headers=['Agent', 'Reads From', 'Writes To'],
    rows=[
        ['orchestrator',         'agent_heartbeats, agent_messages',                                              'agent_heartbeats (via db.heartbeat)'],
        ['google_sync',          '(Google Ads API)',                                                              'campaigns, performance_metrics'],
        ['meta_sync',            '(Meta API)',                                                                    'campaigns, ad_groups, performance_metrics'],
        ['microsoft_sync',       '(Microsoft Ads API)',                                                           'campaigns, ad_groups, performance_metrics'],
        ['performance_analyst',  'campaigns, ad_groups, performance_metrics, keywords',                          'budget_snapshots, search_terms, performance_metrics'],
        ['meta_manager',         'campaigns, ad_groups, performance_metrics',                                    'campaigns, ad_groups, optimization_actions, agent_heartbeats'],
        ['microsoft_manager',    'campaigns, ad_groups, performance_metrics, optimization_actions',             'campaigns, ad_groups, optimization_actions, agent_heartbeats'],
        ['guardrails',           'performance_metrics, campaigns, agent_heartbeats',                            'guardrail_violations'],
        ['bid_budget_optimizer', 'campaigns, keywords, agent_heartbeats',                                       'campaigns, keywords, optimization_actions, agent_heartbeats'],
        ['creative_agent',       'ads, creative_briefs, performance_metrics',                                   'creative_briefs, ads, agent_heartbeats'],
        ['reporting_agent',      'performance_metrics, optimization_actions, budget_snapshots, agent_messages, creative_briefs, negative_keywords', 'reports'],
        ['supabase_sync',        'ALL tables',                                                                   '(Supabase via HTTP POST)'],
    ],
    col_widths=[1.5, 2.5, 2.6],
)

# ---------------------------------------------------------------------------
# Section 4 — Orchestration & Triggers
# ---------------------------------------------------------------------------
heading('4. Orchestration & Triggers')
doc.add_paragraph()

add_table(
    headers=['Trigger', 'What Runs'],
    rows=[
        ['schedule.every().day.at("08:00")',       'Full pipeline (Phases 0–5)'],
        ['schedule.every(5).minutes',              'process_messages() — drains inter-agent message queue'],
        ['schedule.every().hour',                  'System status log'],
        ['python orchestrator.py --once',          'Single pipeline run, no daemon'],
        ['python orchestrator.py --daemon',        'Continuous daemon with cron loop'],
        ['python <agent>.py  (any agent)',         'Standalone run for debugging or manual trigger'],
    ],
    col_widths=[2.8, 3.8],
)

doc.add_paragraph()
body('Active day logic: Phases 2 (bid/budget) and 3 (creative) only execute on Friday, Saturday, Sunday, and Monday. Phases 0, 1, 1b, 4, and 5 run every day regardless of day of week.')

# ---------------------------------------------------------------------------
# Section 5 — Inter-Agent Messaging
# ---------------------------------------------------------------------------
heading('5. Inter-Agent Messaging  (agent_messages table)')
body('Agents communicate via the agent_messages table. The daemon loop drains it every 5 minutes using FOR UPDATE SKIP LOCKED for safe concurrent access.')
doc.add_paragraph()

add_table(
    headers=['from_agent', 'to_agent', 'message_type', 'Purpose'],
    rows=[
        ['performance_analyst',  'orchestrator', 'performance_analysis_complete', 'Signals analysis done, passes summary metrics'],
        ['performance_analyst',  'orchestrator', 'critical_alert',                'Triggers immediate escalation review'],
        ['bid_budget_optimizer', 'all',          'optimization_complete',         'Broadcasts bid/budget run results'],
        ['meta_manager',         'orchestrator', 'optimization_complete',         'Signals Meta optimization pass done'],
        ['microsoft_manager',    'orchestrator', 'optimization_complete',         'Signals Microsoft optimization pass done'],
        ['creative_agent',       'orchestrator', 'creative_work_complete',        'Signals creative refresh done'],
        ['creative_agent',       'orchestrator', 'creative_brief_ready',          'New brief available for review'],
        ['creative_agent',       'orchestrator', 'ads_refreshed',                 'Confirms new ads pushed to platform'],
        ['reporting_agent',      'orchestrator', 'report_complete',               'Daily report written to DB'],
        ['reporting_agent',      'orchestrator', 'report_ready',                  'Report ready for dashboard display'],
    ],
    col_widths=[1.5, 1.2, 1.8, 2.2],
)

# ---------------------------------------------------------------------------
# Section 6 — Guardrails
# ---------------------------------------------------------------------------
heading('6. Guardrails')
body('File: guardrails.py   |   Agent name: guardrails')
body('guardrails.check_pipeline_state() runs between Phase 1b and Phase 2 on every pipeline execution.')
doc.add_paragraph()

heading('Thresholds', level=2)
add_table(
    headers=['Check', 'Threshold', 'Type', 'Effect'],
    rows=[
        ['Daily spend (managed campaigns)',         '> $150/day',    'Escalation', 'Halts Phases 2 & 3, sends email alert'],
        ['7-day CPL',                               '> $1,000',      'Escalation', 'Halts Phases 2 & 3, sends email alert'],
        ['Single keyword spend with 0 conversions', '> $75',         'Escalation', 'Halts Phases 2 & 3, sends email alert'],
        ['API failure duration',                    '> 2 hours',     'Escalation', 'Halts Phases 2 & 3, sends email alert'],
        ['Daily underspend',                        '< 70% of $125', 'Alert',      'Logged only, pipeline continues'],
        ['Single campaign budget increase',         '> 20% per cycle','Block',     'Action rejected, not executed'],
        ['Total budget reallocation per day',       '> $37.50',      'Block',      'Action rejected, not executed'],
        ['Keyword pauses per campaign per day',     '> 3',           'Block',      'Action rejected, not executed'],
        ['Max bid increase per cycle',              '> 25%',         'Block',      'Action rejected, not executed'],
        ['Max bid decrease per cycle',              '> 30%',         'Block',      'Action rejected, not executed'],
        ['Min active keywords per campaign',        '< 5',           'Block',      'Action rejected, not executed'],
    ],
    col_widths=[2.2, 1.4, 0.9, 2.1],
)

heading('Escalation Flow', level=2)
escalation_steps = [
    '1. guardrails.check_pipeline_state() runs between Phase 1b and Phase 2',
    '2. Any escalation violation → written to guardrail_violations table',
    '3. orchestrator sets optimization_halted = True',
    '4. Phases 2 & 3 are skipped for the entire pipeline run',
    '5. Email alert sent to henry@ridgecrestdesigns.com via Resend API',
    '6. Violation marked escalation_sent = true in guardrail_violations',
    '7. Per-action filtering: guardrails.filter_actions(proposed_actions) pre-screens',
    '   individual bid/budget actions — blocked actions are logged but never executed',
]
for step in escalation_steps:
    code_block(step)

doc.add_paragraph()

heading('Tables Used by Guardrails', level=2)
add_table(
    headers=['Table', 'Operation', 'Purpose'],
    rows=[
        ['performance_metrics',  'READ',  'Check today\'s spend and 7-day CPL'],
        ['campaigns',            'READ',  'Filter to claude_code managed campaigns only'],
        ['agent_heartbeats',     'READ',  'Detect API failure duration (last_success_at)'],
        ['guardrail_violations', 'WRITE', 'Log each violation with escalation_sent flag'],
    ],
    col_widths=[1.8, 1.0, 3.8],
)

# ---------------------------------------------------------------------------
# Footer
# ---------------------------------------------------------------------------
doc.add_paragraph()
footer = doc.add_paragraph('Generated by Claude Code  |  Ridgecrest Designs Marketing Automation  |  March 2026')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.runs[0].font.size = Pt(8)
footer.runs[0].font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ---------------------------------------------------------------------------
# Save
# ---------------------------------------------------------------------------
path = '/root/agent/Ridgecrest_Agent_Architecture.docx'
doc.save(path)
print(f'Saved: {path}')
